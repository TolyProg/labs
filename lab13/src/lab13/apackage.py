import Shape, Parallelepiped, Tetrahedron, Sphere
shape_dict = {
    'Parallelepiped': Parallelepiped.Parallelepiped,
    'Tetrahedron': Tetrahedron.Tetrahedron,
    'Sphere': Sphere.Sphere
}

def list_shapes():
    return list(shape_dict.keys())
def list_materials():
    return list(Shape.Shape.material_list.keys())

class Empty(Shape.Shape):
    name = 'Not selected'
    def volume(self):  return 'Not selected'
    def surface(self): return 'Not selected'
    def mass(self):    return 'Not selected'
selected = Empty(0,0,0,0)

def select(x: str, a, b, c, m):
    global selected
    try:
        selected = shape_dict[x](a,b,c,m)
        return f'Selected "{x}"'
    except:
        selected = Empty(0,0,0,0)
        return f'Unknown shape "{x}"'

from docx import Document
def report():
    a = selected.a
    b = selected.b
    c = selected.c
    m = selected.m
    v = selected.volume()
    s = selected.surface()
    d = Document()
    d.add_heading('Отчёт', 0)
    d.add_paragraph(f'Фигура: {selected.name}')
    d.add_paragraph(f'a: {a}')
    d.add_paragraph(f'b: {b}')
    d.add_paragraph(f'c: {c}')
    d.add_paragraph(f's: {s}')
    d.add_paragraph(f'v: {v}')
    d.add_paragraph(f'm: {m}')
    d.save('report.docx')
    return 'OK'
