import par, tet, bal

def list_shapes():
    return ['Parallelepiped', 'Tetrahedron', 'Sphere']
def list_materials():
    return list(material_list.keys())

class Empty:
    name = 'Not selected'
    def volume(a,b,c,d): return 'Not selected'
    def surface(a,b,c,d): return 'Not selected'
selected = Empty()

def select(x: str):
    global selected, volume, surface
    match x:
        case 'Parallelepiped':
            selected = par
        case 'Tetrahedron':
            selected = tet
        case 'Sphere':
            selected = bal
        case _:
            return f'Unknown shape "{x}"'
    volume = selected.volume
    surface = selected.surface

material_list = {'Steel': 7850, 'Glass': 2500, 'Concrete': 2000}

def mass(a,b,c,m):
    try: mm = material_list[m]
    except: return f'Unknown material "{m}"'
    return selected.volume(a,b,c)*mm

from docx import Document
def report(a,b,c,m):
    m = mass(a,b,c,m)
    v = selected.volume(a,b,c)
    s = selected.surface(a,b,c)
    d = Document()
    d.add_heading('Отчёт', 0)
    d.add_paragraph(f'Фигура: {selected.name}')
    d.add_paragraph(f'a: {a}')
    d.add_paragraph(f'b: {b}')
    d.add_paragraph(f'c: {c}')
    d.add_paragraph(f's: {s}')
    d.add_paragraph(f'v: {v}')
    d.add_paragraph(f'm: {m}')
    d.save('report.docx')
    return 'OK'
